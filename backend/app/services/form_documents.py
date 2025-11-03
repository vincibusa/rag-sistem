from __future__ import annotations

import io
import re
from typing import Any, Dict, Iterable, List, Optional, Sequence
from uuid import UUID, uuid4

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from fastapi import UploadFile
from sqlalchemy.exc import OperationalError
from sqlalchemy.orm import Session

from app.core.exceptions import AppException
from app.core.logging import logger
from app.models import FormDocument, FormField as FormFieldModel
from app.schemas.document import AutoFillRequest, AutoFillResponse, FormField
from app.services.rag import RagRetrievalService

from .form_agents import (
    DocumentCompletionAgent,
    PlaceholderDetectionAgent,
    PlaceholderDescriptor,
    RagQueryAgent,
)


class FormDocumentService:
    """Service per la gestione dei documenti form e l'auto-compilazione tramite agenti AI."""

    def __init__(self, session: Session):
        self.session = session
        self.rag_service = RagRetrievalService()
        self._field_counter = 0
        self._registered_field_keys: set[str] = set()
        self._placeholder_agent: PlaceholderDetectionAgent | None = None
        try:
            self._placeholder_agent = PlaceholderDetectionAgent()
        except Exception as exc:  # pragma: no cover - dipendenza esterna
            logger.warning(
                "Impossibile inizializzare il PlaceholderDetectionAgent (%s). Verrà usato il fallback regex.",
                exc,
            )
        self._query_agent = RagQueryAgent()
        self._completion_agent = DocumentCompletionAgent()

    async def upload_form_document(self, file: UploadFile) -> FormDocument:
        """Carica un documento form senza processarlo nel sistema RAG."""
        try:
            file_data = await file.read()
            form_type = self._detect_form_type(file.filename, file.content_type)
            form_document = FormDocument(
                id=uuid4(),
                filename=file.filename or "unnamed",
                content_type=file.content_type or "application/octet-stream",
                data=file_data,
                form_type=form_type,
                size_bytes=len(file_data),
            )
            self.session.add(form_document)
            self.session.commit()
            logger.info("Documento form caricato: %s (%s)", form_document.id, form_type)
            return form_document
        except Exception as exc:
            self.session.rollback()
            logger.error("Errore durante l'upload del documento form: %s", exc)
            raise AppException(f"Errore durante l'upload del documento form: {exc}") from exc

    def extract_form_fields(self, form_id: UUID) -> List[FormField]:
        """Estrae tutti i campi da un documento form."""
        form_document = self._get_form_document(form_id)
        self._registered_field_keys.clear()
        self._field_counter = 0

        try:
            if form_document.form_type == "pdf":
                fields = self._extract_pdf_form_fields(form_document)
            elif form_document.form_type == "word":
                fields = self._extract_word_form_fields(form_document)
            else:
                raise AppException(f"Tipo di form non supportato: {form_document.form_type}")

            self._save_form_fields(form_id, fields)
            logger.info("Estratti %s campi dal documento form %s", len(fields), form_id)
            return fields
        except Exception as exc:
            logger.error("Errore durante l'estrazione dei campi dal form %s: %s", form_id, exc)
            raise AppException(f"Errore durante l'estrazione dei campi: {exc}") from exc

    def auto_fill_form(self, form_id: UUID, request: AutoFillRequest) -> AutoFillResponse:
        """Auto-compila un documento form usando una squadra di agenti AI e il sistema RAG."""
        form_document = self._get_form_document(form_id)
        fields = self._get_form_fields(form_id)

        if not fields:
            raise AppException("Nessun campo trovato nel documento form")

        logger.info("Inizio auto-compilazione per form %s con %s campi", form_id, len(fields))
        logger.info("Contesto di ricerca fornito: %s", request.search_context)
        logger.info("Istruzioni agente: %s", request.agent_guidance)

        if request.field_names:
            fields = [field for field in fields if field.name in request.field_names]

        filled_fields: List[FormField] = []
        search_queries: List[str] = []
        total_confidence = 0.0
        query_cache: dict[str, Sequence[Any]] = {}
        combined_guidance = " ".join(
            part.strip()
            for part in (
                request.search_context or "",
                request.agent_guidance or "",
            )
            if part and part.strip()
        )
        if not combined_guidance:
            combined_guidance = "Compila automaticamente tutti i campi del form."

        for field in fields:
            field_payload = field.model_dump()
            plan = self._query_agent.build_query(field_payload, user_context=combined_guidance or None)
            query = plan.query.strip() or field.name
            search_queries.append(query)

            try:
                logger.info("Ricerca RAG per campo '%s' con query: %s", field.name, query)
                if query in query_cache:
                    rag_results = list(query_cache[query])
                else:
                    rag_results = list(self.rag_service.semantic_search(query=query, top_k=2))
                    query_cache[query] = rag_results
                result_payload = [self._chunk_to_payload(chunk) for chunk in rag_results]
                decision = self._completion_agent.decide(
                    field=field_payload,
                    query=query,
                    chunks=result_payload,
                    guidance=combined_guidance,
                )

                selected_value = (decision.value or "").strip()
                if not selected_value and rag_results:
                    selected_value = self._extract_chunk_text(rag_results[0]).strip()

                field.value = selected_value or field.value
                field.confidence_score = self._combine_confidence(
                    decision=decision,
                    rag_results=rag_results,
                )
                total_confidence += field.confidence_score or 0.0

                logger.info(
                    "Campo '%s' completato con valore '%s' (confidenza %.3f)",
                    field.name,
                    field.value,
                    field.confidence_score,
                )
                if decision.reason:
                    logger.debug("  Motivazione agente: %s", decision.reason)
            except Exception as exc:
                logger.error(
                    "Errore durante la compilazione del campo '%s' (query '%s'): %s",
                    field.name,
                    query,
                    exc,
                )
                field.confidence_score = 0.0

            filled_fields.append(field)

        average_confidence = total_confidence / len(filled_fields) if filled_fields else 0.0
        self._persist_filled_values(form_id, filled_fields)
        compiled_text = self._render_filled_text(form_document, filled_fields)

        return AutoFillResponse(
            form_id=form_id,
            filled_fields=filled_fields,
            total_filled=len([f for f in filled_fields if f.value]),
            average_confidence=average_confidence,
            search_queries=search_queries,
            filled_document_text=compiled_text,
        )

    def get_filled_form(self, form_id: UUID) -> bytes:
        """Genera il documento form compilato."""
        form_document = self._get_form_document(form_id)
        fields = self._get_form_fields(form_id)
        filled_count = sum(1 for field in fields if field.value)
        logger.info(
            "Generazione documento compilato per form %s: %s campi totali, %s con valore.",
            form_id,
            len(fields),
            filled_count,
        )

        try:
            if form_document.form_type == "pdf":
                pdf_bytes = self._fill_pdf_form(form_document, fields)
                logger.info(
                    "Documento PDF compilato generato per form %s (size=%s bytes).",
                    form_id,
                    len(pdf_bytes),
                )
                return pdf_bytes
            if form_document.form_type == "word":
                word_bytes = self._fill_word_form(form_document, fields)
                logger.info(
                    "Documento Word compilato generato per form %s (size=%s bytes).",
                    form_id,
                    len(word_bytes),
                )
                return word_bytes
            raise AppException(f"Tipo di form non supportato per il download: {form_document.form_type}")
        except Exception as exc:
            logger.error("Errore durante la generazione del form compilato %s: %s", form_id, exc)
            raise AppException(f"Errore durante la generazione del documento compilato: {exc}") from exc

    # ---------------------------------------------------------------------
    # Estrazione form e placeholder
    # ---------------------------------------------------------------------

    def _detect_form_type(self, filename: Optional[str], content_type: Optional[str]) -> str:
        if filename and filename.lower().endswith(".pdf"):
            return "pdf"
        if filename and filename.lower().endswith((".docx", ".doc")):
            return "word"
        if content_type:
            if "pdf" in content_type:
                return "pdf"
            if "word" in content_type or "officedocument" in content_type:
                return "word"
        return "unknown"

    def _extract_pdf_form_fields(self, form_document: FormDocument) -> List[FormField]:
        fields: List[FormField] = []

        with fitz.open(stream=form_document.data, filetype="pdf") as doc:
            logger.info("Analisi PDF con %s pagine", len(doc))
            for page_num, page in enumerate(doc):
                widgets = list(page.widgets())
                logger.info("Pagina %s: trovati %s AcroForm widgets", page_num + 1, len(widgets))
                for widget in widgets:
                    form_field = FormField(
                        name=widget.field_name or f"field_{len(fields)}",
                        field_type=self._map_pdf_field_type(widget.field_type),
                        value=widget.field_value,
                        placeholder=widget.field_label,
                        required=bool(widget.field_flags & 1),
                        position={
                            "page": page_num + 1,
                            "x": widget.rect.x0,
                            "y": widget.rect.y0,
                            "width": widget.rect.width,
                            "height": widget.rect.height,
                        },
                        context=f"Pagina {page_num + 1}: {widget.field_label or widget.field_name}",
                    )
                    fields.append(form_field)
                    self._register_field_name(form_field.name)

                text_fields = self._extract_text_placeholders(page, page_num)
                logger.info("Pagina %s: trovati %s placeholder testuali", page_num + 1, len(text_fields))
                fields.extend(text_fields)

        acroform_count = sum(
            1
            for field in fields
            if isinstance(field.position, dict)
            and {"x", "y", "width", "height"}.issubset(field.position.keys())
        )
        logger.info(
            "Totale campi estratti: %s (AcroForm: %s, Placeholder testuali: %s)",
            len(fields),
            acroform_count,
            len(fields) - acroform_count,
        )
        return fields

    def _extract_text_placeholders(self, page, page_num: int) -> List[FormField]:
        if self._placeholder_agent:
            try:
                ai_fields = self._extract_text_placeholders_with_agent(page, page_num)
                if ai_fields:
                    return ai_fields
            except Exception as exc:  # pragma: no cover - dipendenza esterna
                logger.warning(
                    "Analisi AI per placeholder fallita alla pagina %s: %s. Fallback a regex.",
                    page_num + 1,
                    exc,
                )
        return self._extract_text_placeholders_with_regex(page, page_num)

    def _extract_text_placeholders_with_agent(self, page, page_num: int) -> List[FormField]:
        text = page.get_text()
        logger.debug("Analisi placeholder AI su pagina %s (prime 500 chars): %s", page_num + 1, text[:500])
        descriptors = self._placeholder_agent.analyse(text, page_num + 1) if self._placeholder_agent else []
        if not descriptors:
            return []
        return self._convert_ai_fields_to_form_fields(descriptors, page_num, page, text)

    def _extract_text_placeholders_with_regex(self, page, page_num: int) -> List[FormField]:
        fields: List[FormField] = []
        try:
            text = page.get_text()
            logger.debug("Fallback regex su pagina %s (prime 500 chars): %s", page_num + 1, text[:500])
            patterns = [
                r"_{5,}",
                r"\s_{3,}\s",
                r"\(_{2,}\)",
                r"\.{3,}",
                r"-{3,}",
                r"\s{10,}",
            ]
            for pattern in patterns:
                for match in re.finditer(pattern, text):
                    placeholder_text = match.group()
                    line_start = text.rfind("\n", 0, match.start()) + 1
                    line_end = text.find("\n", match.end())
                    line_end = len(text) if line_end == -1 else line_end
                    full_line = text[line_start:line_end].strip()
                    context_start = max(line_start, match.start() - 80)
                    context_text = text[context_start:match.start()].strip()
                    field_name = self._generate_field_name_from_context(full_line, context_text)
                    bbox = None
                    trimmed_placeholder = placeholder_text.strip()
                    if trimmed_placeholder:
                        try:
                            rects = page.search_for(trimmed_placeholder, hit_max=1)
                        except Exception:
                            rects = []
                        if rects:
                            rect = rects[0]
                            bbox = [rect.x0, rect.y0, rect.x1, rect.y1]
                    position = {
                        "page": page_num + 1,
                        "text_position": match.start(),
                    }
                    if bbox:
                        position["bbox"] = bbox

                    form_field = FormField(
                        name=field_name,
                        field_type="text",
                        value=None,
                        placeholder=placeholder_text,
                        required=False,
                        position=position,
                        context=full_line or f"Pagina {page_num + 1}: campo da compilare",
                    )
                    fields.append(form_field)
                    self._register_field_name(form_field.name)
        except Exception as exc:  # pragma: no cover - dipendenza esterna
            logger.warning(
                "Errore durante l'estrazione regex dei placeholder alla pagina %s: %s",
                page_num + 1,
                exc,
            )
        return fields

    def _convert_ai_fields_to_form_fields(
        self,
        descriptors: Iterable[PlaceholderDescriptor],
        page_num: int,
        page,
        page_text: str,
    ) -> List[FormField]:
        fields: List[FormField] = []
        search_offset = 0

        for descriptor in descriptors:
            placeholder = (descriptor.placeholder_text or "").strip()
            context = descriptor.context.strip()
            name = self._ensure_unique_field_name(descriptor.name)
            text_position = None
            bbox = None

            if placeholder:
                idx = page_text.find(placeholder, search_offset)
                if idx == -1:
                    idx = page_text.find(placeholder)
                if idx != -1:
                    text_position = idx
                    search_offset = idx + len(placeholder)
                try:
                    rects = page.search_for(placeholder, hit_max=1)
                except Exception:
                    rects = []
                if rects:
                    rect = rects[0]
                    bbox = [rect.x0, rect.y0, rect.x1, rect.y1]

            metadata = {
                "type": descriptor.type,
                "query": descriptor.query,
                "raw_name": descriptor.name,
                "placeholder_text": descriptor.placeholder_text,
                "context": descriptor.context,
            }
            position: Dict[str, Any] = {"page": page_num + 1, "ai": metadata}
            if text_position is not None:
                position["text_position"] = text_position
            if bbox:
                position["bbox"] = bbox

            field = FormField(
                name=name,
                field_type=self._map_ai_field_type(descriptor.type),
                value=None,
                placeholder=placeholder or None,
                required=False,
                position=position,
                context=context or f"Pagina {page_num + 1}: campo da compilare",
            )
            fields.append(field)
            self._register_field_name(field.name)

            logger.debug(
                "Placeholder AI '%s' -> campo '%s' (tipo %s, query '%s')",
                placeholder,
                field.name,
                descriptor.type,
                descriptor.query,
            )

        return fields

    # ---------------------------------------------------------------------
    # Supporto RAG e persistenza
    # ---------------------------------------------------------------------

    def _chunk_to_payload(self, chunk: Any) -> dict[str, Any]:
        metadata = getattr(chunk, "metadata", {}) or {}
        return {
            "id": getattr(chunk, "id", None),
            "text": self._extract_chunk_text(chunk),
            "score": self._extract_chunk_score(chunk),
            "metadata": self._to_dict(metadata),
        }

    def _extract_chunk_text(self, chunk: Any) -> str:
        for attr in ("text", "content", "value"):
            value = getattr(chunk, attr, None)
            if isinstance(value, str) and value.strip():
                return value
        if isinstance(chunk, dict):
            for key in ("text", "content", "value"):
                value = chunk.get(key)
                if isinstance(value, str) and value.strip():
                    return value
        metadata = getattr(chunk, "metadata", None)
        if isinstance(metadata, dict):
            for key in ("text", "content", "value"):
                value = metadata.get(key)
                if isinstance(value, str) and value.strip():
                    return value
        payload = getattr(chunk, "payload", None)
        if isinstance(payload, dict):
            for key in ("text", "content"):
                value = payload.get(key)
                if isinstance(value, str) and value.strip():
                    return value
        return ""

    def _extract_chunk_score(self, chunk: Any) -> float:
        score = getattr(chunk, "score", None)
        if score is not None:
            try:
                return float(score)
            except (ValueError, TypeError):
                pass

        distance = getattr(chunk, "distance", None)
        if distance is not None:
            try:
                distance_val = float(distance)
            except (ValueError, TypeError):
                distance_val = None
            if distance_val is not None:
                return max(0.0, 1.0 - distance_val)

        metadata = getattr(chunk, "metadata", None)
        if isinstance(metadata, dict):
            for key in ("score", "similarity", "confidence"):
                value = metadata.get(key)
                if value is not None:
                    try:
                        return float(value)
                    except (ValueError, TypeError):
                        continue
        if isinstance(chunk, dict):
            for key in ("score", "similarity", "confidence"):
                value = chunk.get(key)
                if value is not None:
                    try:
                        return float(value)
                    except (ValueError, TypeError):
                        continue
        return 0.0

    def _combine_confidence(self, *, decision, rag_results: Sequence[Any]) -> float:
        confidence = decision.confidence if hasattr(decision, "confidence") else 0.0
        if decision.selected_chunk_index is not None and 0 <= decision.selected_chunk_index < len(rag_results):
            chunk_conf = self._extract_chunk_score(rag_results[decision.selected_chunk_index])
            confidence = max(confidence, chunk_conf)
        return round(confidence, 3)

    def _persist_filled_values(self, form_id: UUID, fields: Iterable[FormField]) -> None:
        def _apply_updates() -> None:
            model_map = {
                model.name: model
                for model in self.session.query(FormFieldModel).filter(
                    FormFieldModel.form_document_id == form_id
                )
            }
            for field in fields:
                model = model_map.get(field.name)
                if not model:
                    continue
                model.value = field.value
                model.confidence_score = field.confidence_score
                model.context = field.context
                model.placeholder = field.placeholder
                model.position = field.position
            self.session.commit()

        try:
            _apply_updates()
        except OperationalError as exc:
            logger.warning(
                "Connessione al database interrotta durante il salvataggio dei campi del form %s. Ritento una volta.",
                form_id,
            )
            self.session.rollback()
            _apply_updates()

    # ---------------------------------------------------------------------
    # Utilità per naming e mapping
    # ---------------------------------------------------------------------

    def _generate_field_name_from_context(self, full_line: str, context_text: str) -> str:
        field_labels = [
            "nato a",
            "residente a",
            "via/piazza",
            "impresa",
            "sede legale",
            "sede operativa",
            "codice fiscale",
            "partita iva",
            "e-mail",
            "pec",
            "nome",
            "cognome",
            "data di nascita",
            "indirizzo",
            "città",
            "provincia",
            "cap",
            "telefono",
            "fax",
            "sito web",
            "ragione sociale",
            "forma giuridica",
        ]
        for label in field_labels:
            if label in full_line.lower():
                field_name = label.replace(" ", "_").replace("/", "_")
                logger.debug("Trovata etichetta campo: '%s' -> '%s'", label, field_name)
                return self._ensure_unique_field_name(field_name)

        if context_text:
            cleaned = re.sub(r"[^a-zA-Z0-9\s]", " ", context_text)
            words = cleaned.strip().split()
            if len(words) >= 2:
                name_words = words[-3:] if len(words) >= 3 else words[-2:]
                field_name = "_".join(name_words).lower()
                if len(field_name) >= 3:
                    return self._ensure_unique_field_name(field_name)

        return self._next_generic_field_name()

    def _normalize_field_key(self, name: str | None) -> str:
        if not name:
            return ""
        return re.sub(r"[^a-zA-Z0-9_]+", "_", name.strip()).strip("_").lower()

    def _register_field_name(self, name: str | None) -> None:
        key = self._normalize_field_key(name)
        if key:
            self._registered_field_keys.add(key)

    def _ensure_unique_field_name(self, candidate: str | None) -> str:
        base_key = self._normalize_field_key(candidate)
        if not base_key:
            return self._next_generic_field_name()
        unique_key = base_key
        suffix = 1
        while unique_key in self._registered_field_keys:
            suffix += 1
            unique_key = f"{base_key}_{suffix}"
        self._registered_field_keys.add(unique_key)
        return unique_key

    def _next_generic_field_name(self) -> str:
        self._field_counter += 1
        generic = f"field_{self._field_counter}"
        self._registered_field_keys.add(generic)
        return generic

    def _map_pdf_field_type(self, pdf_field_type: int) -> str:
        return {
            1: "button",
            2: "checkbox",
            3: "radio",
            4: "text",
            5: "select",
            6: "signature",
        }.get(pdf_field_type, "unknown")

    def _map_ai_field_type(self, ai_type: str) -> str:
        mapping = {
            "data_nascita": "date",
            "data_di_nascita": "date",
            "data": "date",
            "telefono": "tel",
            "cellulare": "tel",
            "email": "email",
            "pec": "email",
            "codice_fiscale": "text",
            "partita_iva": "text",
            "numero": "number",
            "quantita": "number",
            "firma": "signature",
        }
        return mapping.get((ai_type or "").lower(), "text")

    # ---------------------------------------------------------------------
    # Persistence helpers
    # ---------------------------------------------------------------------

    def _get_form_document(self, form_id: UUID) -> FormDocument:
        form_document = self.session.get(FormDocument, form_id)
        if not form_document:
            raise AppException(f"Documento form {form_id} non trovato")
        return form_document

    def _get_form_fields(self, form_id: UUID) -> List[FormField]:
        field_models = (
            self.session.query(FormFieldModel)
            .filter(FormFieldModel.form_document_id == form_id)
            .all()
        )
        return [
            FormField(
                name=model.name,
                field_type=model.field_type,
                value=model.value,
                placeholder=model.placeholder,
                required=model.required,
                position=model.position,
                context=model.context,
                confidence_score=model.confidence_score,
            )
            for model in field_models
        ]

    def _save_form_fields(self, form_id: UUID, fields: Iterable[FormField]) -> None:
        self.session.query(FormFieldModel).filter(
            FormFieldModel.form_document_id == form_id
        ).delete()

        for field in fields:
            model = FormFieldModel(
                id=uuid4(),
                form_document_id=form_id,
                name=field.name,
                field_type=field.field_type,
                value=field.value,
                placeholder=field.placeholder,
                required=field.required,
                position=field.position,
                context=field.context,
                confidence_score=field.confidence_score,
            )
            self.session.add(model)
        self.session.commit()

    # ---------------------------------------------------------------------
    # Compilazione documenti
    # ---------------------------------------------------------------------

    def _fill_pdf_form(self, form_document: FormDocument, fields: List[FormField]) -> bytes:
        with fitz.open(stream=form_document.data, filetype="pdf") as doc:
            total_widgets_filled = 0
            placeholder_overlays = 0
            for page_num in range(len(doc)):
                page = doc[page_num]
                widget_field_names: set[str] = set()
                widgets = list(page.widgets())

                for widget in widgets:
                    matching_field = next(
                        (f for f in fields if f.name == widget.field_name and f.value),
                        None,
                    )
                    if matching_field:
                        widget.field_value = matching_field.value
                        widget.update()
                        if widget.field_name:
                            widget_field_names.add(widget.field_name)
                        total_widgets_filled += 1

                page_fields = [
                    field
                    for field in fields
                    if field.value
                    and isinstance(field.position, dict)
                    and field.position.get("page") == page_num + 1
                    and field.name not in widget_field_names
                ]

                for field in page_fields:
                    position = field.position or {}
                    bbox = position.get("bbox") if isinstance(position, dict) else None
                    if not bbox:
                        continue
                    try:
                        rect = fitz.Rect(*bbox)
                    except Exception:
                        continue
                    rect = fitz.Rect(rect.x0 - 1, rect.y0 - 1, rect.x1 + 1, rect.y1 + 1)
                    fontsize = max(6, min(14, rect.height * 0.75))
                    page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                    page.insert_textbox(
                        rect,
                        field.value or "",
                        fontname="helv",
                        fontsize=fontsize,
                        color=(0, 0, 0),
                        align=0,
                    )
                    placeholder_overlays += 1

            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            logger.info(
                "Compilazione PDF completata: %s widgets , %s placeholder testuali aggiornati.",
                total_widgets_filled,
                placeholder_overlays,
            )
            return output_buffer.getvalue()

    def _fill_word_form(self, form_document: FormDocument, fields: List[FormField]) -> bytes:
        return form_document.data

    def _render_filled_text(self, form_document: FormDocument, fields: List[FormField]) -> str | None:
        if not fields:
            return None

        try:
            if form_document.form_type == "pdf":
                return self._render_pdf_text(form_document, fields)
            if form_document.form_type == "word":
                return self._render_word_text(form_document, fields)
        except Exception as exc:  # pragma: no cover - dipendenza esterna
            logger.warning(
                "Impossibile generare testo compilato per form %s: %s",
                form_document.id,
                exc,
            )

        return self._render_fallback_summary(fields)

    def _render_pdf_text(self, form_document: FormDocument, fields: List[FormField]) -> str:
        page_outputs: List[str] = []
        with fitz.open(stream=form_document.data, filetype="pdf") as doc:
            for page_num, page in enumerate(doc):
                text = page.get_text()
                page_fields = [
                    field
                    for field in fields
                    if field.value
                    and isinstance(field.position, dict)
                    and field.position.get("page") == page_num + 1
                ]
                replacements: List[tuple[int | None, str, str]] = []
                for field in page_fields:
                    placeholder = field.placeholder or ""
                    value = field.value or ""
                    position = field.position or {}
                    text_pos = position.get("text_position")
                    if text_pos is not None and placeholder:
                        replacements.append((text_pos, placeholder, value))
                    elif placeholder:
                        replacements.append((None, placeholder, value))
                replacements.sort(key=lambda item: item[0] if item[0] is not None else -1, reverse=True)
                for text_pos, placeholder, value in replacements:
                    if text_pos is not None and placeholder:
                        segment = text[text_pos : text_pos + len(placeholder)]
                        if segment == placeholder:
                            text = text[:text_pos] + value + text[text_pos + len(placeholder) :]
                        else:
                            text = text.replace(placeholder, value, 1)
                    elif placeholder:
                        text = text.replace(placeholder, value, 1)
                page_outputs.append(text)
        return "\n\n".join(page_outputs)

    def _render_word_text(self, form_document: FormDocument, fields: List[FormField]) -> str:
        buffer = io.BytesIO(form_document.data)
        document = DocxDocument(buffer)
        for paragraph in document.paragraphs:
            for field in fields:
                if not field.value or not field.placeholder:
                    continue
                if field.placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(field.placeholder, field.value)
        output = io.StringIO()
        for paragraph in document.paragraphs:
            output.write(paragraph.text)
            output.write("\n")
        return output.getvalue()

    def _render_fallback_summary(self, fields: List[FormField]) -> str:
        lines = []
        for field in fields:
            if not field.value:
                continue
            lines.append(f"{field.name}: {field.value}")
        return "\n".join(lines) if lines else ""

    # ---------------------------------------------------------------------
    # Utils
    # ---------------------------------------------------------------------

    @staticmethod
    def _to_dict(metadata: Any) -> dict[str, Any]:
        if metadata is None:
            return {}
        if isinstance(metadata, dict):
            return metadata
        return {"value": str(metadata)}
